import io
import os
import datetime
from django.conf import settings
from django.db.models import Count, Q
from django.http import HttpResponse
from django.utils import timezone
from rest_framework import generics, status, permissions, filters
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.views import APIView

from docx import Document
from docx.shared import Pt
from .models import Appeal, AppealFile, AppealHistory, Comment
from .serializers import (
    AppealListSerializer, AppealDetailSerializer,
    PublicAppealCreateSerializer, CommentSerializer, AppealHistorySerializer,
)
from .emails import (
    send_appeal_registered, send_status_changed,
    send_executor_assigned, send_appeal_closed,
)


class IsOperatorOrAdmin(permissions.BasePermission):
    def has_permission(self, request, view):
        return request.user.is_authenticated and (
            request.user.role in ['operator', 'admin'] or request.user.is_superuser
        )


class IsExecutorOrHigher(permissions.BasePermission):
    def has_permission(self, request, view):
        return request.user.is_authenticated


# ──────────────────────────────────────────────────────────────────────────────
# Публичные эндпоинты (без авторизации)
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([permissions.AllowAny])
def public_create_appeal(request):
    """POST /api/appeals/public/ — гражданин подаёт обращение."""
    serializer = PublicAppealCreateSerializer(data=request.data)
    if serializer.is_valid():
        appeal = serializer.save()

        # Прикрепляем файлы (до 5)
        files = request.FILES.getlist('files')[:5]
        for f in files:
            AppealFile.objects.create(
                appeal=appeal,
                file=f,
                original_name=f.name,
            )

        # Лог создания
        AppealHistory.objects.create(
            appeal=appeal,
            action=f'Обращение зарегистрировано через публичную форму',
        )

        # Email-уведомление гражданину
        send_appeal_registered(appeal)

        return Response({
            'registration_number': appeal.registration_number,
            'message': 'Ваше обращение зарегистрировано. Ответ будет направлен в течение 10 рабочих дней.',
        }, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def public_check_status(request):
    """GET /api/appeals/check-status/?number=2026-0042."""
    number = request.query_params.get('number')
    if not number:
        return Response({'error': 'Укажите номер обращения'}, status=400)
    try:
        appeal = Appeal.objects.get(registration_number=number)
        return Response({
            'registration_number': appeal.registration_number,
            'status': appeal.status,
            'status_display': appeal.get_status_display(),
            'created_at': appeal.created_at,
            'subject': appeal.subject,
            'category': appeal.get_category_display(),
        })
    except Appeal.DoesNotExist:
        return Response({'error': 'Обращение не найдено'}, status=404)


# ──────────────────────────────────────────────────────────────────────────────
# Список и детальный просмотр обращений
# ──────────────────────────────────────────────────────────────────────────────

class AppealListCreateView(generics.ListCreateAPIView):
    """GET /api/appeals/ — список (оператор/админ); POST — создание (staff)."""
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['registration_number', 'citizen_full_name']
    ordering_fields = ['created_at', 'deadline', 'status']
    ordering = ['-created_at']

    def get_serializer_class(self):
        if self.request.method == 'POST':
            return AppealDetailSerializer
        return AppealListSerializer

    def get_queryset(self):
        user = self.request.user
        qs = Appeal.objects.select_related('assigned_to')

        # Исполнитель видит только свои заявки
        if user.role == 'executor' and not user.is_superuser:
            qs = qs.filter(assigned_to=user)

        # Фильтры
        status_filter = self.request.query_params.get('status')
        category_filter = self.request.query_params.get('category')
        executor_filter = self.request.query_params.get('assigned_to')
        date_from = self.request.query_params.get('date_from')
        date_to = self.request.query_params.get('date_to')
        overdue = self.request.query_params.get('overdue')

        if status_filter:
            qs = qs.filter(status=status_filter)
        if category_filter:
            qs = qs.filter(category=category_filter)
        if executor_filter:
            qs = qs.filter(assigned_to_id=executor_filter)
        if date_from:
            qs = qs.filter(created_at__date__gte=date_from)
        if date_to:
            qs = qs.filter(created_at__date__lte=date_to)
        if overdue == 'true':
            today = datetime.date.today()
            qs = qs.filter(
                deadline__lt=today
            ).exclude(status__in=['done', 'closed'])

        return qs


class AppealDetailView(generics.RetrieveUpdateAPIView):
    """GET/PATCH /api/appeals/{id}/."""
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = AppealDetailSerializer
    queryset = Appeal.objects.all()

    def perform_update(self, serializer):
        old = self.get_object()
        old_status = old.status
        old_executor = old.assigned_to_id

        appeal = serializer.save()
        user = self.request.user

        # Логируем изменение статуса
        if old_status != appeal.status:
            AppealHistory.objects.create(
                appeal=appeal,
                user=user,
                action=f'Статус изменён с «{old.get_status_display()}» на «{appeal.get_status_display()}»',
            )
            # Email-уведомление при смене статуса
            send_status_changed(appeal, old.get_status_display(), appeal.get_status_display())
            # Дополнительное уведомление при закрытии
            if appeal.status == Appeal.Status.CLOSED:
                send_appeal_closed(appeal)

        # Логируем смену исполнителя
        if old_executor != appeal.assigned_to_id:
            executor_name = appeal.assigned_to.full_name if appeal.assigned_to else 'не назначен'
            AppealHistory.objects.create(
                appeal=appeal,
                user=user,
                action=f'Исполнитель назначен: {executor_name}',
            )
            # Email исполнителю при назначении
            if appeal.assigned_to:
                send_executor_assigned(appeal)


# ──────────────────────────────────────────────────────────────────────────────
# Комментарии
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
def appeal_comments(request, pk):
    try:
        appeal = Appeal.objects.get(pk=pk)
    except Appeal.DoesNotExist:
        return Response({'error': 'Не найдено'}, status=404)

    if request.method == 'GET':
        comments = appeal.comments.select_related('author')
        return Response(CommentSerializer(comments, many=True).data)

    elif request.method == 'POST':
        text = request.data.get('text', '').strip()
        if not text:
            return Response({'error': 'Текст комментария не может быть пустым'}, status=400)
        comment = Comment.objects.create(appeal=appeal, author=request.user, text=text)
        # Лог
        AppealHistory.objects.create(
            appeal=appeal,
            user=request.user,
            action=f'Добавлен комментарий: «{text[:80]}{"..." if len(text) > 80 else ""}»',
        )
        return Response(CommentSerializer(comment).data, status=201)


# ──────────────────────────────────────────────────────────────────────────────
# История
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def appeal_history(request, pk):
    try:
        appeal = Appeal.objects.get(pk=pk)
    except Appeal.DoesNotExist:
        return Response({'error': 'Не найдено'}, status=404)
    history = appeal.history.select_related('user')
    return Response(AppealHistorySerializer(history, many=True).data)


# ──────────────────────────────────────────────────────────────────────────────
# Загрузка файлов
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def upload_files(request, pk):
    try:
        appeal = Appeal.objects.get(pk=pk)
    except Appeal.DoesNotExist:
        return Response({'error': 'Не найдено'}, status=404)

    existing_count = appeal.files.count()
    files = request.FILES.getlist('files')
    slots = 5 - existing_count
    if slots <= 0:
        return Response({'error': 'Максимум 5 файлов на обращение'}, status=400)

    created = []
    for f in files[:slots]:
        af = AppealFile.objects.create(appeal=appeal, file=f, original_name=f.name)
        created.append({'id': af.id, 'name': af.original_name})

    return Response({'uploaded': created}, status=201)


# ──────────────────────────────────────────────────────────────────────────────
# Генерация Word-ответа
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_response(request, pk):
    try:
        appeal = Appeal.objects.get(pk=pk)
    except Appeal.DoesNotExist:
        return Response({'error': 'Не найдено'}, status=404)

    if not appeal.resolution_text:
        return Response({'error': 'Сначала заполните текст решения'}, status=400)

    # Создаём документ
    doc = Document()

    # Стили
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Заголовок
    heading = doc.add_paragraph()
    heading_run = heading.add_run('Ответ на обращение гражданина')
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading.alignment = 1  # CENTER

    doc.add_paragraph()

    # Обращение
    doc.add_paragraph(f'Уважаемый(-ая) {appeal.citizen_full_name}!')
    doc.add_paragraph()

    # Основной текст
    intro = doc.add_paragraph()
    intro.add_run(
        f'На Ваше обращение № {appeal.registration_number} '
        f'от {appeal.created_at.strftime("%d.%m.%Y")} сообщаем:'
    )
    doc.add_paragraph()

    # Текст решения
    doc.add_paragraph(appeal.resolution_text)
    doc.add_paragraph()

    # Подпись
    executor_name = appeal.assigned_to.full_name if appeal.assigned_to else '________________'
    doc.add_paragraph(f'С уважением,')
    doc.add_paragraph(executor_name)

    # Дата
    doc.add_paragraph()
    doc.add_paragraph(datetime.date.today().strftime('%d.%m.%Y'))

    # Сохраняем
    responses_dir = settings.RESPONSES_ROOT
    os.makedirs(responses_dir, exist_ok=True)
    filename = f"response_{appeal.registration_number}_{datetime.date.today()}.docx"
    filepath = os.path.join(responses_dir, filename)
    doc.save(filepath)

    # Лог
    AppealHistory.objects.create(
        appeal=appeal,
        user=request.user,
        action=f'Сгенерирован официальный ответ: {filename}',
    )

    # Отдаём файл
    from django.http import FileResponse
    return FileResponse(
        open(filepath, 'rb'),
        as_attachment=True,
        filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


# ──────────────────────────────────────────────────────────────────────────────
# Дашборд
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def dashboard(request):
    today = datetime.date.today()
    month_start = today.replace(day=1)
    week_ago = today - datetime.timedelta(days=7)

    total = Appeal.objects.count()
    in_progress = Appeal.objects.filter(status__in=['assigned', 'in_progress', 'on_site']).count()
    overdue = Appeal.objects.filter(
        deadline__lt=today
    ).exclude(status__in=['done', 'closed']).count()
    done_month = Appeal.objects.filter(
        status='done',
        updated_at__date__gte=month_start
    ).count()

    # По категориям
    by_category = list(
        Appeal.objects.values('category').annotate(count=Count('id'))
    )
    category_labels = dict(Appeal.Category.choices)
    by_category_labeled = [
        {'category': item['category'], 'label': category_labels.get(item['category'], item['category']), 'count': item['count']}
        for item in by_category
    ]

    # Последние 5
    recent = Appeal.objects.select_related('assigned_to').order_by('-created_at')[:5]
    recent_data = AppealListSerializer(recent, many=True).data

    # Просроченные (для таблицы)
    overdue_list = Appeal.objects.filter(
        deadline__lt=today
    ).exclude(status__in=['done', 'closed']).select_related('assigned_to').order_by('deadline')[:10]
    overdue_data = AppealListSerializer(overdue_list, many=True).data

    # Рейтинг исполнителей за неделю
    from apps.users.models import User
    executor_rating = list(
        Appeal.objects.filter(
            status='done',
            updated_at__date__gte=week_ago,
            assigned_to__isnull=False,
        ).values('assigned_to__full_name', 'assigned_to__id').annotate(
            closed=Count('id')
        ).order_by('-closed')[:5]
    )

    return Response({
        'total': total,
        'in_progress': in_progress,
        'overdue': overdue,
        'done_month': done_month,
        'by_category': by_category_labeled,
        'recent': recent_data,
        'overdue_list': overdue_data,
        'executor_rating': executor_rating,
    })


# ──────────────────────────────────────────────────────────────────────────────
# Экспорт реестра (Excel / PDF)
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def export_appeals(request):
    print("!!! INSIDE export_appeals view !!!")
    """
    GET /api/appeals/export/?format=excel  — выгрузка в .xlsx
    GET /api/appeals/export/?format=pdf   — выгрузка в .pdf
    Поддерживает те же query-параметры фильтрации, что и /api/appeals/.
    """
    fmt = request.query_params.get('format', 'excel').lower()

    # Применяем те же фильтры
    qs = Appeal.objects.select_related('assigned_to').order_by('-created_at')
    status_filter = request.query_params.get('status')
    category_filter = request.query_params.get('category')
    executor_filter = request.query_params.get('assigned_to')
    date_from = request.query_params.get('date_from')
    date_to = request.query_params.get('date_to')
    search = request.query_params.get('search')
    if status_filter:
        qs = qs.filter(status=status_filter)
    if category_filter:
        qs = qs.filter(category=category_filter)
    if executor_filter:
        qs = qs.filter(assigned_to_id=executor_filter)
    if date_from:
        qs = qs.filter(created_at__date__gte=date_from)
    if date_to:
        qs = qs.filter(created_at__date__lte=date_to)
    if search:
        qs = qs.filter(
            Q(registration_number__icontains=search) |
            Q(citizen_full_name__icontains=search)
        )

    appeals = list(qs[:5000])  # ограничение на выгрузку

    if fmt == 'excel':
        return _export_excel(appeals)
    elif fmt == 'pdf':
        return _export_pdf(appeals)
    else:
        return Response({'error': 'Поддерживаемые форматы: excel, pdf'}, status=400)

export_appeals.view_class.format_kwarg = None


def _export_excel(appeals):
    """Генерирует .xlsx и возвращает как FileResponse."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Реестр обращений'

    # Заголовки
    headers = [
        '№', 'Рег. номер', 'Дата', 'ФИО гражданина', 'Телефон',
        'Адрес', 'Категория', 'Тема', 'Статус', 'Исполнитель',
        'Срок', 'Просрочено', 'Дата обновления',
    ]
    header_fill = PatternFill('solid', fgColor='1F3A5F')
    header_font = Font(color='FFFFFF', bold=True, size=11)

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    ws.row_dimensions[1].height = 30

    # Ширины колонок
    col_widths = [5, 15, 12, 25, 16, 30, 15, 35, 15, 25, 12, 12, 20]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    status_display = dict(Appeal.Status.choices)
    category_display = dict(Appeal.Category.choices)

    # Данные
    for row_idx, appeal in enumerate(appeals, 2):
        executor = appeal.assigned_to.full_name if appeal.assigned_to else '—'
        deadline_str = appeal.deadline.strftime('%d.%m.%Y') if appeal.deadline else '—'
        created_str = appeal.created_at.strftime('%d.%m.%Y %H:%M')
        updated_str = appeal.updated_at.strftime('%d.%m.%Y %H:%M')
        overdue_str = 'Да' if appeal.is_overdue else 'Нет'

        row_data = [
            row_idx - 1,
            appeal.registration_number,
            created_str,
            appeal.citizen_full_name,
            appeal.citizen_phone,
            appeal.citizen_address,
            category_display.get(appeal.category, appeal.category),
            appeal.subject,
            status_display.get(appeal.status, appeal.status),
            executor,
            deadline_str,
            overdue_str,
            updated_str,
        ]
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(vertical='top', wrap_text=True)
            # Выделяем просроченные красным
            if appeal.is_overdue:
                cell.fill = PatternFill('solid', fgColor='FFF0F0')

        ws.row_dimensions[row_idx].height = 20

    # Автофильтр
    ws.auto_filter.ref = f'A1:{openpyxl.utils.get_column_letter(len(headers))}1'

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    filename = f'appeals_export_{datetime.date.today()}.xlsx'
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _export_pdf(appeals):
    """Генерирует простой PDF-отчёт через reportlab."""
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return Response(
            {'error': 'reportlab не установлен. Выполните: pip install reportlab'},
            status=500,
        )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title', parent=styles['Heading1'],
        fontSize=14, spaceAfter=6,
    )
    small_style = ParagraphStyle('Small', parent=styles['Normal'], fontSize=8)

    elements = []
    elements.append(Paragraph('Реестр обращений граждан', title_style))
    elements.append(Paragraph(
        f'Дата выгрузки: {datetime.date.today().strftime("%d.%m.%Y")} · Всего записей: {len(appeals)}',
        small_style,
    ))
    elements.append(Spacer(1, 6 * mm))

    # Таблица
    status_display = dict(Appeal.Status.choices)
    category_display = dict(Appeal.Category.choices)

    table_data = [[
        '№', 'Номер', 'Дата', 'ФИО гражданина', 'Категория', 'Тема', 'Статус', 'Исполнитель', 'Срок'
    ]]

    for i, appeal in enumerate(appeals, 1):
        executor = appeal.assigned_to.full_name if appeal.assigned_to else '—'
        deadline_str = appeal.deadline.strftime('%d.%m.%Y') if appeal.deadline else '—'
        table_data.append([
            str(i),
            appeal.registration_number,
            appeal.created_at.strftime('%d.%m.%Y'),
            Paragraph(appeal.citizen_full_name, small_style),
            category_display.get(appeal.category, appeal.category),
            Paragraph(appeal.subject[:60], small_style),
            status_display.get(appeal.status, appeal.status),
            executor,
            deadline_str,
        ])

    col_widths = [8*mm, 22*mm, 18*mm, 40*mm, 22*mm, 65*mm, 22*mm, 35*mm, 18*mm]
    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F3A5F')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F8FF')]),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#CCCCCC')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    elements.append(tbl)

    doc.build(elements)
    buffer.seek(0)

    filename = f'appeals_export_{datetime.date.today()}.pdf'
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ──────────────────────────────────────────────────────────────────────────────
# Личный кабинет гражданина — публичный поиск обращений
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def citizen_portal(request):
    """
    GET /api/appeals/citizen/?email=user@example.com
    GET /api/appeals/citizen/?number=2026-0001
    GET /api/appeals/citizen/?email=user@example.com&number=2026-0001

    Публичный эндпоинт: позволяет гражданину получить список своих обращений
    по email или регистрационному номеру.
    """
    email = request.query_params.get('email', '').strip()
    number = request.query_params.get('number', '').strip()

    if not email and not number:
        return Response(
            {'error': 'Укажите email или номер обращения'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    qs = Appeal.objects.all()

    if email and number:
        qs = qs.filter(citizen_email__iexact=email, registration_number=number)
    elif email:
        qs = qs.filter(citizen_email__iexact=email)
    else:
        qs = qs.filter(registration_number=number)

    if not qs.exists():
        return Response({'results': [], 'message': 'Обращения не найдены'}, status=200)

    data = []
    for appeal in qs.order_by('-created_at')[:50]:
        data.append({
            'registration_number': appeal.registration_number,
            'subject': appeal.subject,
            'category': appeal.get_category_display(),
            'status': appeal.status,
            'status_display': appeal.get_status_display(),
            'created_at': appeal.created_at,
            'updated_at': appeal.updated_at,
            'deadline': appeal.deadline,
            'is_overdue': appeal.is_overdue,
            'resolution_text': appeal.resolution_text or None,
        })

    return Response({'results': data})


# ──────────────────────────────────────────────────────────────────────────────
# Личный кабинет исполнителя
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def my_appeals(request):
    """GET /api/my-appeals/ — поручения текущего исполнителя."""
    user = request.user
    qs = Appeal.objects.filter(assigned_to=user).order_by('-created_at')
    data = AppealListSerializer(qs, many=True).data
    return Response(data)


@api_view(['PATCH'])
@permission_classes([permissions.IsAuthenticated])
def my_appeal_update_status(request, pk):
    """PATCH /api/my-appeals/<pk>/update-status/ — смена статуса исполнителем."""
    try:
        appeal = Appeal.objects.get(pk=pk, assigned_to=request.user)
    except Appeal.DoesNotExist:
        return Response({'error': 'Не найдено или нет доступа'}, status=404)

    new_status = request.data.get('status')
    allowed = [Appeal.Status.IN_PROGRESS, Appeal.Status.ON_SITE, Appeal.Status.DONE]
    if new_status not in [s.value for s in allowed]:
        return Response(
            {'error': f'Допустимые статусы: {[s.value for s in allowed]}'},
            status=400,
        )

    old_display = appeal.get_status_display()
    appeal.status = new_status
    appeal.save(update_fields=['status', 'updated_at'])

    AppealHistory.objects.create(
        appeal=appeal,
        user=request.user,
        action=f'Статус изменён с «{old_display}» на «{appeal.get_status_display()}»',
    )
    send_status_changed(appeal, old_display, appeal.get_status_display())

    return Response(AppealListSerializer(appeal).data)


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def list_executors(request):
    """GET /api/users/executors/ — список исполнителей для выпадающих списков."""
    from apps.users.models import User
    executors = User.objects.filter(role='executor').values('id', 'full_name', 'department')
    return Response(list(executors))
